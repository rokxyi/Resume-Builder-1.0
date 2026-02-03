import os
from pathlib import Path
from typing import Optional
from docx import Document
from pypdf import PdfReader
import aiofiles


class DocumentParser:
    """Service for parsing different document formats"""
    
    @staticmethod
    async def parse_docx(file_path: str) -> str:
        """Extract text from .docx file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
    
    @staticmethod
    async def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text_content = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            
            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
    
    @staticmethod
    async def parse_txt(file_path: str) -> str:
        """Extract text from .txt file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except Exception as e:
            raise ValueError(f"Error parsing TXT file: {str(e)}")
    
    @staticmethod
    async def parse_file(file_path: str, file_type: str) -> str:
        """Parse file based on type"""
        if file_type == "application/pdf":
            return await DocumentParser.parse_pdf(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            return await DocumentParser.parse_docx(file_path)
        elif file_type == "text/plain":
            return await DocumentParser.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
