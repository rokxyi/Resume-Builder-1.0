import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import re


class ResumeGenerator:
    """Service for generating formatted .docx resumes"""
    
    @staticmethod
    def _parse_llm_response(raw_response: str) -> Dict[str, Any]:
        """Parse LLM response to extract JSON data"""
        try:
            # Try to find JSON in the response
            json_match = re.search(r'\{.*\}', raw_response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                return json.loads(json_str)
            else:
                # If no JSON found, return raw response
                return {"error": "Could not parse JSON from response", "raw": raw_response}
        except json.JSONDecodeError:
            return {"error": "Invalid JSON in response", "raw": raw_response}
    
    @staticmethod
    def generate_docx(resume_data: Dict[str, Any], output_path: str) -> str:
        """Generate a formatted .docx resume"""
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        resume = resume_data.get("resume", {})
        
        # Header - Name
        name = resume.get("name", "Candidate Name")
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_run.font.name = 'Arial'
        
        # Contact Info
        contact = resume.get("contact", {})
        contact_parts = []
        if contact.get("location"):
            contact_parts.append(contact["location"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" • ".join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # Professional Summary
        if resume.get("professional_summary"):
            summary_heading = doc.add_paragraph()
            summary_heading_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
            summary_heading_run.font.size = Pt(11)
            summary_heading_run.font.bold = True
            summary_heading_run.font.name = 'Arial'
            
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(resume["professional_summary"])
            summary_run.font.size = Pt(10)
            summary_run.font.name = 'Arial'
            doc.add_paragraph()
        
        # Core Competencies
        if resume.get("core_competencies"):
            comp_heading = doc.add_paragraph()
            comp_heading_run = comp_heading.add_run("CORE COMPETENCIES")
            comp_heading_run.font.size = Pt(11)
            comp_heading_run.font.bold = True
            comp_heading_run.font.name = 'Arial'
            
            for category, skills in resume["core_competencies"].items():
                comp_para = doc.add_paragraph()
                comp_para.paragraph_format.left_indent = Inches(0.25)
                
                category_run = comp_para.add_run(f"{category}: ")
                category_run.font.size = Pt(10)
                category_run.font.bold = True
                category_run.font.name = 'Arial'
                
                skills_run = comp_para.add_run(" | ".join(skills))
                skills_run.font.size = Pt(10)
                skills_run.font.name = 'Arial'
            
            doc.add_paragraph()
        
        # Professional Experience
        if resume.get("experience"):
            exp_heading = doc.add_paragraph()
            exp_heading_run = exp_heading.add_run("PROFESSIONAL EXPERIENCE")
            exp_heading_run.font.size = Pt(11)
            exp_heading_run.font.bold = True
            exp_heading_run.font.name = 'Arial'
            
            for exp in resume["experience"]:
                # Company and location
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(f"{exp.get('company', '')} • {exp.get('location', '')}")
                company_run.font.size = Pt(10)
                company_run.font.bold = True
                company_run.font.name = 'Arial'
                
                # Title and dates
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{exp.get('title', '')} | {exp.get('start_date', '')} - {exp.get('end_date', '')}")
                title_run.font.size = Pt(10)
                title_run.font.bold = True
                title_run.font.name = 'Arial'
                
                # Bullets
                for bullet in exp.get("bullets", []):
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    for run in bullet_para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                
                doc.add_paragraph()
        
        # Education
        if resume.get("education"):
            edu_heading = doc.add_paragraph()
            edu_heading_run = edu_heading.add_run("EDUCATION")
            edu_heading_run.font.size = Pt(11)
            edu_heading_run.font.bold = True
            edu_heading_run.font.name = 'Arial'
            
            for edu in resume["education"]:
                edu_para = doc.add_paragraph()
                
                degree_text = f"{edu.get('degree', '')} ({edu.get('field', '')})" if edu.get('field') else edu.get('degree', '')
                degree_run = edu_para.add_run(degree_text)
                degree_run.font.size = Pt(10)
                degree_run.font.bold = True
                degree_run.font.name = 'Arial'
                
                edu_para.add_run("\n")
                
                uni_parts = [edu.get('university', '')]
                if edu.get('location'):
                    uni_parts.append(edu['location'])
                if edu.get('gpa'):
                    uni_parts.append(f"GPA: {edu['gpa']}")
                if edu.get('graduation_date'):
                    uni_parts.append(edu['graduation_date'])
                
                uni_run = edu_para.add_run(" • ".join(uni_parts))
                uni_run.font.size = Pt(10)
                uni_run.font.name = 'Arial'
                
                if edu.get('coursework'):
                    coursework_para = doc.add_paragraph()
                    coursework_para.paragraph_format.left_indent = Inches(0.25)
                    coursework_run = coursework_para.add_run(f"Relevant Coursework: {', '.join(edu['coursework'])}")
                    coursework_run.font.size = Pt(10)
                    coursework_run.font.name = 'Arial'
            
            doc.add_paragraph()
        
        # Certifications
        if resume.get("certifications"):
            cert_heading = doc.add_paragraph()
            cert_heading_run = cert_heading.add_run("CERTIFICATIONS")
            cert_heading_run.font.size = Pt(11)
            cert_heading_run.font.bold = True
            cert_heading_run.font.name = 'Arial'
            
            for cert in resume["certifications"]:
                cert_para = doc.add_paragraph(cert, style='List Bullet')
                cert_para.paragraph_format.left_indent = Inches(0.25)
                for run in cert_para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
        
        # Save document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return output_path
